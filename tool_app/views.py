import os
import tempfile
from io import BytesIO
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse, FileResponse
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
import json
from docx import Document
from docx.shared import Inches

from .models import FileConversion, Newsletter
from .forms import (
    FileUploadForm, TextToPdfForm, ImageCompressionForm, ImageConversionForm, QRCodeForm,
    MetaTagForm, URLEncoderDecoderForm, DomainResolverForm, WhoisLookupForm, RobotsSitemapForm,
    HashGeneratorForm, JWTDecoderForm, SSLCheckerForm, EmailValidatorForm, TextEncryptionForm,
    MemeGeneratorForm, EmojiTranslatorForm, RandomQuoteForm, RandomNameForm, UnitConverterForm,
    AudioConverterForm, AudioSpeedChangerForm, VideoToAudioForm, NewsletterForm
)
from PIL import Image, ImageOps
import qrcode
from qrcode.image.svg import SvgImage
import urllib.parse
import socket
import dns.resolver
import whois
from datetime import datetime
import xml.etree.ElementTree as ET
import hashlib
import base64
import json
import re
import ssl
from datetime import datetime, timezone
from email.utils import parseaddr
import random
import string
from PIL import Image, ImageDraw, ImageFont
import textwrap
import os
import tempfile
from pydub import AudioSegment
try:
    from moviepy.editor import VideoFileClip
except ImportError:
    VideoFileClip = None


def home(request):
    """Home page with tool overview"""
    return render(request, 'tool_app/home.html')


@csrf_exempt
@require_POST
def newsletter_subscribe(request):
    """Handle newsletter subscription"""
    try:
        if request.content_type == 'application/json':
            # Handle JSON request (AJAX)
            data = json.loads(request.body)
            email = data.get('email')
        else:
            # Handle form submission
            email = request.POST.get('email')
        
        if not email:
            return JsonResponse({'error': 'Email is required'}, status=400)
        
        # Get client info
        ip_address = request.META.get('HTTP_X_FORWARDED_FOR')
        if ip_address:
            ip_address = ip_address.split(',')[0].strip()
        else:
            ip_address = request.META.get('REMOTE_ADDR')
        
        user_agent = request.META.get('HTTP_USER_AGENT', '')
        
        # Check if email already exists and is active
        existing = Newsletter.objects.filter(email=email).first()
        
        if existing:
            if existing.is_active:
                return JsonResponse({
                    'error': 'This email is already subscribed to our newsletter.'
                }, status=400)
            else:
                # Reactivate existing subscription
                existing.is_active = True
                existing.unsubscribed_at = None
                existing.ip_address = ip_address
                existing.user_agent = user_agent
                existing.save()
                
                return JsonResponse({
                    'success': True,
                    'message': 'Successfully resubscribed to newsletter!'
                })
        else:
            # Create new subscription
            Newsletter.objects.create(
                email=email,
                ip_address=ip_address,
                user_agent=user_agent
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Successfully subscribed to newsletter!'
            })
            
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'An error occurred: {str(e)}'}, status=500)


def file_converter(request):
    """File converter tool page"""
    if request.method == 'POST':
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            conversion = form.save(commit=False)
            conversion.status = 'processing'
            conversion.save()
            
            # Process the conversion
            try:
                converted_file = process_file_conversion(conversion)
                if converted_file:
                    conversion.converted_file = converted_file
                    conversion.status = 'completed'
                    conversion.save()
                    messages.success(request, 'File converted successfully!')
                    return redirect('tool_app:conversion_result', pk=conversion.pk)
                else:
                    conversion.status = 'failed'
                    conversion.error_message = 'Conversion failed'
                    conversion.save()
                    messages.error(request, 'File conversion failed. Please try again.')
            except Exception as e:
                conversion.status = 'failed'
                conversion.error_message = str(e)
                conversion.save()
                messages.error(request, f'Error during conversion: {str(e)}')
    else:
        form = FileUploadForm()
    
    return render(request, 'tool_app/file_converter.html', {'form': form})


def text_to_pdf(request):
    """Text to PDF converter"""
    if request.method == 'POST':
        form = TextToPdfForm(request.POST)
        if form.is_valid():
            try:
                pdf_buffer = create_pdf_from_text(
                    form.cleaned_data['text_content'],
                    form.cleaned_data.get('title', 'Document')
                )
                
                response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
                response['Content-Disposition'] = 'attachment; filename="converted_document.pdf"'
                return response
            except Exception as e:
                messages.error(request, f'Error creating PDF: {str(e)}')
    else:
        form = TextToPdfForm()
    
    return render(request, 'tool_app/text_to_pdf.html', {'form': form})


def conversion_result(request, pk):
    """Display conversion result"""
    conversion = get_object_or_404(FileConversion, pk=pk)
    return render(request, 'tool_app/conversion_result.html', {'conversion': conversion})


def download_file(request, pk):
    """Download converted file"""
    conversion = get_object_or_404(FileConversion, pk=pk)
    if conversion.converted_file:
        response = FileResponse(
            conversion.converted_file.open(),
            as_attachment=True,
            filename=conversion.converted_filename
        )
        return response
    else:
        messages.error(request, 'File not found or conversion not completed.')
        return redirect('tool_app:home')


# JSON API Views
@csrf_exempt
@require_POST
def api_convert_file(request):
    """API endpoint for file conversion"""
    try:
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file provided'}, status=400)
        
        file = request.FILES['file']
        conversion_type = request.POST.get('conversion_type')
        
        if not conversion_type:
            return JsonResponse({'error': 'No conversion type specified'}, status=400)
        
        # Create conversion record
        conversion = FileConversion.objects.create(
            original_file=file,
            conversion_type=conversion_type,
            status='processing'
        )
        
        # Process conversion
        converted_file = process_file_conversion(conversion)
        
        if converted_file:
            conversion.converted_file = converted_file
            conversion.status = 'completed'
            conversion.save()
            
            return JsonResponse({
                'success': True,
                'conversion_id': conversion.id,
                'download_url': f'/download/{conversion.id}/',
                'original_filename': conversion.original_filename,
                'converted_filename': conversion.converted_filename
            })
        else:
            conversion.status = 'failed'
            conversion.save()
            return JsonResponse({'error': 'Conversion failed'}, status=500)
            
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_POST
def api_text_to_pdf(request):
    """API endpoint for text to PDF conversion"""
    try:
        data = json.loads(request.body)
        text_content = data.get('text_content')
        title = data.get('title', 'Document')
        
        if not text_content:
            return JsonResponse({'error': 'No text content provided'}, status=400)
        
        pdf_buffer = create_pdf_from_text(text_content, title)
        
        # Return base64 encoded PDF
        import base64
        pdf_base64 = base64.b64encode(pdf_buffer.getvalue()).decode('utf-8')
        
        return JsonResponse({
            'success': True,
            'pdf_data': pdf_base64,
            'filename': f'{title.replace(" ", "_")}.pdf'
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def api_conversion_status(request, pk):
    """API endpoint to check conversion status"""
    try:
        conversion = get_object_or_404(FileConversion, pk=pk)
        return JsonResponse({
            'id': conversion.id,
            'status': conversion.status,
            'conversion_type': conversion.conversion_type,
            'original_filename': conversion.original_filename,
            'converted_filename': conversion.converted_filename,
            'created_at': conversion.created_at.isoformat(),
            'error_message': conversion.error_message
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# Helper functions
def process_file_conversion(conversion):
    """Process file conversion based on type"""
    try:
        if conversion.conversion_type == 'txt_to_pdf':
            return convert_txt_to_pdf(conversion.original_file)
        elif conversion.conversion_type == 'pdf_to_txt':
            return convert_pdf_to_txt(conversion.original_file)
        elif conversion.conversion_type == 'doc_to_pdf':
            return convert_doc_to_pdf(conversion.original_file)
        elif conversion.conversion_type == 'pdf_to_doc':
            return convert_pdf_to_doc(conversion.original_file)
        elif conversion.conversion_type == 'image_compress':
            return compress_image_file(conversion.original_file)
        elif conversion.conversion_type == 'image_convert':
            return convert_image_format_file(conversion.original_file)
        else:
            raise ValueError(f"Unsupported conversion type: {conversion.conversion_type}")
    except Exception as e:
        print(f"Conversion error: {e}")
        return None


def convert_txt_to_pdf(txt_file):
    """Convert text file to PDF"""
    try:
        # Read text content
        content = txt_file.read().decode('utf-8')
        
        # Create PDF
        pdf_buffer = create_pdf_from_text(content, "Converted Document")
        
        # Save to file
        filename = f"converted_{txt_file.name.split('.')[0]}.pdf"
        return ContentFile(pdf_buffer.getvalue(), name=filename)
        
    except Exception as e:
        print(f"TXT to PDF conversion error: {e}")
        return None


def convert_pdf_to_txt(pdf_file):
    """Convert PDF to text file"""
    try:
        # This is a simplified implementation
        # For production, you'd use PyPDF2 or pdfplumber
        content = "PDF to text conversion is not fully implemented in this demo.\nThis would require additional libraries like PyPDF2 or pdfplumber."
        
        filename = f"converted_{pdf_file.name.split('.')[0]}.txt"
        return ContentFile(content.encode('utf-8'), name=filename)
        
    except Exception as e:
        print(f"PDF to TXT conversion error: {e}")
        return None


def convert_doc_to_pdf(doc_file):
    """Convert DOC/DOCX to PDF"""
    try:
        # Read DOCX content
        doc = Document(doc_file)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Convert to PDF
        pdf_buffer = create_pdf_from_text(content, "Converted Document")
        
        filename = f"converted_{doc_file.name.split('.')[0]}.pdf"
        return ContentFile(pdf_buffer.getvalue(), name=filename)
        
    except Exception as e:
        print(f"DOC to PDF conversion error: {e}")
        return None


def convert_pdf_to_doc(pdf_file):
    """Convert PDF to DOC/DOCX"""
    try:
        # This is a simplified implementation
        content = "PDF to DOC conversion is not fully implemented in this demo.\nThis would require additional libraries for PDF parsing."
        
        # Create DOCX
        doc = Document()
        doc.add_paragraph(content)
        
        # Save to buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        filename = f"converted_{pdf_file.name.split('.')[0]}.docx"
        return ContentFile(doc_buffer.getvalue(), name=filename)
        
    except Exception as e:
        print(f"PDF to DOC conversion error: {e}")
        return None


def create_pdf_from_text(text_content, title):
    """Create PDF from text content"""
    buffer = BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Add title
    story.append(Paragraph(title, styles['Title']))
    
    # Add content
    for line in text_content.split('\n'):
        if line.strip():
            story.append(Paragraph(line, styles['Normal']))
        else:
            story.append(Paragraph('<br/>', styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


# Image Processing Views
def image_compression(request):
    """Image compression tool"""
    if request.method == 'POST':
        form = ImageCompressionForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                compressed_file = compress_image(
                    form.cleaned_data['image_file'],
                    form.cleaned_data['quality'],
                    form.cleaned_data.get('resize_width')
                )
                
                if compressed_file:
                    original_filename = form.cleaned_data['image_file'].name
                    filename = f"compressed_{original_filename}"
                    
                    response = HttpResponse(compressed_file.getvalue(), content_type='image/jpeg')
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
                else:
                    messages.error(request, 'Image compression failed.')
            except Exception as e:
                messages.error(request, f'Error compressing image: {str(e)}')
    else:
        form = ImageCompressionForm()
    
    return render(request, 'tool_app/image_compression.html', {'form': form})


def image_conversion(request):
    """Image format conversion tool"""
    if request.method == 'POST':
        form = ImageConversionForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                converted_file = convert_image_format(
                    form.cleaned_data['image_file'],
                    form.cleaned_data['target_format'],
                    form.cleaned_data['quality']
                )
                
                if converted_file:
                    original_name = form.cleaned_data['image_file'].name.split('.')[0]
                    target_ext = form.cleaned_data['target_format'].lower()
                    if target_ext == 'jpeg':
                        target_ext = 'jpg'
                    filename = f"{original_name}_converted.{target_ext}"
                    
                    content_type = f"image/{target_ext}"
                    if target_ext == 'svg':
                        content_type = 'image/svg+xml'
                    
                    response = HttpResponse(converted_file.getvalue(), content_type=content_type)
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
                else:
                    messages.error(request, 'Image conversion failed.')
            except Exception as e:
                messages.error(request, f'Error converting image: {str(e)}')
    else:
        form = ImageConversionForm()
    
    return render(request, 'tool_app/image_conversion.html', {'form': form})


def qr_code_generator(request):
    """QR code generation tool"""
    if request.method == 'POST':
        form = QRCodeForm(request.POST)
        if form.is_valid():
            try:
                qr_file = generate_qr_code(
                    form.cleaned_data['content'],
                    form.cleaned_data['size'],
                    form.cleaned_data['format']
                )
                
                if qr_file:
                    format_ext = form.cleaned_data['format'].lower()
                    filename = f"qrcode.{format_ext}"
                    
                    content_type = f"image/{format_ext}"
                    if format_ext == 'svg':
                        content_type = 'image/svg+xml'
                    
                    response = HttpResponse(qr_file.getvalue(), content_type=content_type)
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
                else:
                    messages.error(request, 'QR code generation failed.')
            except Exception as e:
                messages.error(request, f'Error generating QR code: {str(e)}')
    else:
        form = QRCodeForm()
    
    return render(request, 'tool_app/qr_code_generator.html', {'form': form})


# Web & SEO Tools Views
def meta_tag_generator(request):
    """Meta tag generator for SEO"""
    meta_tags = None
    if request.method == 'POST':
        form = MetaTagForm(request.POST)
        if form.is_valid():
            try:
                meta_tags = generate_meta_tags(form.cleaned_data)
            except Exception as e:
                messages.error(request, f'Error generating meta tags: {str(e)}')
    else:
        form = MetaTagForm()
    
    return render(request, 'tool_app/meta_tag_generator.html', {
        'form': form, 
        'meta_tags': meta_tags
    })


def url_encoder_decoder(request):
    """URL encoder/decoder tool"""
    result = None
    if request.method == 'POST':
        form = URLEncoderDecoderForm(request.POST)
        if form.is_valid():
            try:
                url_text = form.cleaned_data['url_text']
                operation = form.cleaned_data['operation']
                
                if operation == 'encode':
                    result = urllib.parse.quote(url_text)
                else:  # decode
                    result = urllib.parse.unquote(url_text)
                    
            except Exception as e:
                messages.error(request, f'Error processing URL: {str(e)}')
    else:
        form = URLEncoderDecoderForm()
    
    return render(request, 'tool_app/url_encoder_decoder.html', {
        'form': form,
        'result': result
    })


def domain_ip_resolver(request):
    """Domain to IP resolver tool"""
    dns_results = None
    if request.method == 'POST':
        form = DomainResolverForm(request.POST)
        if form.is_valid():
            try:
                domain = form.cleaned_data['domain']
                record_type = form.cleaned_data['record_type']
                dns_results = resolve_dns_records(domain, record_type)
                
            except Exception as e:
                messages.error(request, f'Error resolving domain: {str(e)}')
    else:
        form = DomainResolverForm()
    
    return render(request, 'tool_app/domain_ip_resolver.html', {
        'form': form,
        'dns_results': dns_results
    })


def whois_lookup(request):
    """Whois lookup tool"""
    whois_data = None
    if request.method == 'POST':
        form = WhoisLookupForm(request.POST)
        if form.is_valid():
            try:
                domain = form.cleaned_data['domain']
                whois_data = get_whois_data(domain)
                
            except Exception as e:
                messages.error(request, f'Error performing whois lookup: {str(e)}')
    else:
        form = WhoisLookupForm()
    
    return render(request, 'tool_app/whois_lookup.html', {
        'form': form,
        'whois_data': whois_data
    })


def robots_sitemap_generator(request):
    """Robots.txt and sitemap.xml generator"""
    generated_files = None
    if request.method == 'POST':
        form = RobotsSitemapForm(request.POST)
        if form.is_valid():
            try:
                generated_files = generate_robots_sitemap(form.cleaned_data)
                
            except Exception as e:
                messages.error(request, f'Error generating files: {str(e)}')
    else:
        form = RobotsSitemapForm()
    
    return render(request, 'tool_app/robots_sitemap_generator.html', {
        'form': form,
        'generated_files': generated_files
    })


# Image Processing Helper Functions
def compress_image(image_file, quality, resize_width=None):
    """Compress image with optional resizing"""
    try:
        # Open image
        image = Image.open(image_file)
        
        # Convert to RGB if necessary (for JPEG compatibility)
        if image.mode in ('RGBA', 'P'):
            image = image.convert('RGB')
        
        # Resize if width specified
        if resize_width:
            width_percent = (resize_width / float(image.size[0]))
            height = int((float(image.size[1]) * float(width_percent)))
            image = image.resize((resize_width, height), Image.Resampling.LANCZOS)
        
        # Save compressed image
        buffer = BytesIO()
        image.save(buffer, format='JPEG', quality=quality, optimize=True)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Image compression error: {e}")
        return None


def convert_image_format(image_file, target_format, quality=85):
    """Convert image to different format"""
    try:
        # Open image
        image = Image.open(image_file)
        
        # Handle different target formats
        buffer = BytesIO()
        
        if target_format == 'JPEG':
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')
            image.save(buffer, format='JPEG', quality=quality, optimize=True)
        elif target_format == 'PNG':
            image.save(buffer, format='PNG', optimize=True)
        elif target_format == 'WEBP':
            image.save(buffer, format='WEBP', quality=quality, optimize=True)
        elif target_format == 'BMP':
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')
            image.save(buffer, format='BMP')
        elif target_format == 'TIFF':
            image.save(buffer, format='TIFF')
        else:
            raise ValueError(f"Unsupported target format: {target_format}")
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Image conversion error: {e}")
        return None


def generate_qr_code(content, size, format_type):
    """Generate QR code"""
    try:
        # Size mapping
        size_map = {
            'small': 200,
            'medium': 400,
            'large': 600,
            'xlarge': 800
        }
        
        pixel_size = size_map.get(size, 400)
        box_size = max(1, pixel_size // 37)  # QR codes are typically 37x37 modules
        
        # Create QR code
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=box_size,
            border=4,
        )
        qr.add_data(content)
        qr.make(fit=True)
        
        buffer = BytesIO()
        
        if format_type == 'SVG':
            # Generate SVG
            img = qr.make_image(image_factory=SvgImage)
            svg_string = img.to_string()
            buffer.write(svg_string)
        else:
            # Generate raster image
            img = qr.make_image(fill_color="black", back_color="white")
            
            if format_type == 'JPEG':
                # Convert to RGB for JPEG
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img if img.mode == 'RGBA' else None)
                rgb_img.save(buffer, format='JPEG', quality=95)
            else:  # PNG
                img.save(buffer, format='PNG')
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"QR code generation error: {e}")
        return None


# File-based processing for API
def compress_image_file(image_file, quality=80):
    """Compress image file with default settings"""
    try:
        compressed = compress_image(image_file, quality)
        if compressed:
            filename = f"compressed_{image_file.name.split('.')[0]}.jpg"
            return ContentFile(compressed.getvalue(), name=filename)
        return None
    except Exception as e:
        print(f"Image file compression error: {e}")
        return None


def convert_image_format_file(image_file, target_format='PNG'):
    """Convert image file format with default settings"""
    try:
        converted = convert_image_format(image_file, target_format)
        if converted:
            ext = target_format.lower()
            if ext == 'jpeg':
                ext = 'jpg'
            filename = f"converted_{image_file.name.split('.')[0]}.{ext}"
            return ContentFile(converted.getvalue(), name=filename)
        return None
    except Exception as e:
        print(f"Image format conversion error: {e}")
        return None


# Web & SEO Helper Functions
def generate_meta_tags(data):
    """Generate HTML meta tags"""
    tags = []
    
    # Basic meta tags
    tags.append(f'<title>{data["title"]}</title>')
    tags.append(f'<meta name="description" content="{data["description"]}">')
    
    if data.get('keywords'):
        tags.append(f'<meta name="keywords" content="{data["keywords"]}">')
    
    if data.get('author'):
        tags.append(f'<meta name="author" content="{data["author"]}">')
    
    # Canonical URL
    if data.get('canonical_url'):
        tags.append(f'<link rel="canonical" href="{data["canonical_url"]}">')
    
    # Open Graph tags
    tags.append(f'<meta property="og:title" content="{data["title"]}">')
    tags.append(f'<meta property="og:description" content="{data["description"]}">')
    tags.append('<meta property="og:type" content="website">')
    
    if data.get('canonical_url'):
        tags.append(f'<meta property="og:url" content="{data["canonical_url"]}">')
    
    if data.get('og_image'):
        tags.append(f'<meta property="og:image" content="{data["og_image"]}">')
    
    # Twitter Card tags
    tags.append('<meta name="twitter:card" content="summary_large_image">')
    tags.append(f'<meta name="twitter:title" content="{data["title"]}">')
    tags.append(f'<meta name="twitter:description" content="{data["description"]}">')
    
    if data.get('og_image'):
        tags.append(f'<meta name="twitter:image" content="{data["og_image"]}">')
    
    # Additional SEO tags
    tags.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    tags.append('<meta charset="UTF-8">')
    tags.append('<meta name="robots" content="index, follow">')
    
    return '\n'.join(tags)


def resolve_dns_records(domain, record_type):
    """Resolve DNS records for a domain"""
    results = []
    
    try:
        # Clean domain name
        domain = domain.strip().lower()
        if domain.startswith(('http://', 'https://')):
            domain = domain.split('//', 1)[1].split('/', 1)[0]
        
        resolver = dns.resolver.Resolver()
        resolver.timeout = 10
        resolver.lifetime = 10
        
        try:
            answers = resolver.resolve(domain, record_type)
            for answer in answers:
                if record_type == 'A':
                    results.append({
                        'type': 'A Record (IPv4)',
                        'value': str(answer.address),
                        'ttl': answer.rrset.ttl
                    })
                elif record_type == 'AAAA':
                    results.append({
                        'type': 'AAAA Record (IPv6)',
                        'value': str(answer.address),
                        'ttl': answer.rrset.ttl
                    })
                elif record_type == 'CNAME':
                    results.append({
                        'type': 'CNAME Record',
                        'value': str(answer.target),
                        'ttl': answer.rrset.ttl
                    })
                elif record_type == 'MX':
                    results.append({
                        'type': 'MX Record',
                        'value': f'{answer.preference} {answer.exchange}',
                        'ttl': answer.rrset.ttl
                    })
                elif record_type == 'NS':
                    results.append({
                        'type': 'NS Record',
                        'value': str(answer.target),
                        'ttl': answer.rrset.ttl
                    })
                elif record_type == 'TXT':
                    results.append({
                        'type': 'TXT Record',
                        'value': str(answer).strip('"'),
                        'ttl': answer.rrset.ttl
                    })
        
        except dns.resolver.NXDOMAIN:
            results.append({
                'type': 'Error',
                'value': f'Domain {domain} does not exist',
                'ttl': None
            })
        except dns.resolver.NoAnswer:
            results.append({
                'type': 'Info',
                'value': f'No {record_type} records found for {domain}',
                'ttl': None
            })
        except Exception as e:
            results.append({
                'type': 'Error',
                'value': f'DNS lookup error: {str(e)}',
                'ttl': None
            })
        
        # Also try to get basic IP info for A records
        if record_type == 'A' and results and results[0]['type'] == 'A Record (IPv4)':
            try:
                ip = results[0]['value']
                hostname = socket.gethostbyaddr(ip)[0]
                results.append({
                    'type': 'Reverse DNS',
                    'value': hostname,
                    'ttl': None
                })
            except:
                pass
        
    except Exception as e:
        results.append({
            'type': 'Error',
            'value': f'DNS resolution failed: {str(e)}',
            'ttl': None
        })
    
    return results


def get_whois_data(domain):
    """Get whois information for a domain"""
    try:
        # Clean domain name
        domain = domain.strip().lower()
        if domain.startswith(('http://', 'https://')):
            domain = domain.split('//', 1)[1].split('/', 1)[0]
        
        w = whois.whois(domain)
        
        # Format the whois data
        whois_info = {
            'domain_name': w.domain_name[0] if isinstance(w.domain_name, list) else w.domain_name,
            'registrar': w.registrar,
            'creation_date': w.creation_date[0] if isinstance(w.creation_date, list) else w.creation_date,
            'expiration_date': w.expiration_date[0] if isinstance(w.expiration_date, list) else w.expiration_date,
            'updated_date': w.updated_date[0] if isinstance(w.updated_date, list) else w.updated_date,
            'status': w.status,
            'name_servers': w.name_servers,
            'registrant_name': getattr(w, 'registrant_name', None),
            'registrant_organization': getattr(w, 'registrant_organization', None),
            'registrant_country': getattr(w, 'registrant_country', None),
            'admin_email': getattr(w, 'admin_email', None),
            'tech_email': getattr(w, 'tech_email', None),
        }
        
        return whois_info
        
    except Exception as e:
        return {
            'error': f'Whois lookup failed: {str(e)}',
            'domain_name': domain
        }


def generate_robots_sitemap(data):
    """Generate robots.txt and sitemap.xml content"""
    domain_url = data['domain_url'].rstrip('/')
    sitemap_urls = [url.strip() for url in data['sitemap_urls'].split('\n') if url.strip()]
    disallow_paths = [path.strip() for path in (data.get('disallow_paths', '') or '').split('\n') if path.strip()]
    crawl_delay = data.get('crawl_delay')
    
    # Generate robots.txt
    robots_content = ['User-agent: *']
    
    if disallow_paths:
        for path in disallow_paths:
            if not path.startswith('/'):
                path = '/' + path
            robots_content.append(f'Disallow: {path}')
    else:
        robots_content.append('Disallow:')
    
    if crawl_delay:
        robots_content.append(f'Crawl-delay: {crawl_delay}')
    
    robots_content.append('')
    robots_content.append(f'Sitemap: {domain_url}/sitemap.xml')
    
    # Generate sitemap.xml
    root = ET.Element('urlset')
    root.set('xmlns', 'http://www.sitemaps.org/schemas/sitemap/0.9')
    
    for url in sitemap_urls:
        if not url.startswith(('http://', 'https://')):
            if url.startswith('/'):
                url = domain_url + url
            else:
                url = domain_url + '/' + url
        
        url_elem = ET.SubElement(root, 'url')
        loc_elem = ET.SubElement(url_elem, 'loc')
        loc_elem.text = url
        
        lastmod_elem = ET.SubElement(url_elem, 'lastmod')
        lastmod_elem.text = datetime.now().strftime('%Y-%m-%d')
        
        changefreq_elem = ET.SubElement(url_elem, 'changefreq')
        changefreq_elem.text = 'monthly'
        
        priority_elem = ET.SubElement(url_elem, 'priority')
        priority_elem.text = '1.0' if url == domain_url or url == domain_url + '/' else '0.8'
    
    # Convert XML to string with proper formatting
    ET.indent(root, space="  ", level=0)
    sitemap_content = '<?xml version="1.0" encoding="UTF-8"?>\n' + ET.tostring(root, encoding='unicode')
    
    return {
        'robots_txt': '\n'.join(robots_content),
        'sitemap_xml': sitemap_content,
        'domain': domain_url
    }


# Security & Utility Tools Views

def hash_generator(request):
    """Hash generator tool"""
    hash_results = None
    if request.method == 'POST':
        form = HashGeneratorForm(request.POST)
        if form.is_valid():
            try:
                text_input = form.cleaned_data['text_input']
                hash_types = form.cleaned_data['hash_types']
                hash_results = generate_hashes(text_input, hash_types)
            except Exception as e:
                messages.error(request, f'Error generating hashes: {str(e)}')
    else:
        form = HashGeneratorForm()
    
    return render(request, 'tool_app/hash_generator.html', {
        'form': form,
        'hash_results': hash_results
    })


def jwt_decoder(request):
    """JWT decoder tool"""
    jwt_data = None
    if request.method == 'POST':
        form = JWTDecoderForm(request.POST)
        if form.is_valid():
            try:
                jwt_token = form.cleaned_data['jwt_token'].strip()
                jwt_data = decode_jwt_token(jwt_token)
            except Exception as e:
                messages.error(request, f'Error decoding JWT: {str(e)}')
    else:
        form = JWTDecoderForm()
    
    return render(request, 'tool_app/jwt_decoder.html', {
        'form': form,
        'jwt_data': jwt_data
    })


def ssl_checker(request):
    """SSL certificate checker tool"""
    ssl_info = None
    if request.method == 'POST':
        form = SSLCheckerForm(request.POST)
        if form.is_valid():
            try:
                domain = form.cleaned_data['domain']
                port = form.cleaned_data['port']
                ssl_info = check_ssl_certificate(domain, port)
            except Exception as e:
                messages.error(request, f'Error checking SSL certificate: {str(e)}')
    else:
        form = SSLCheckerForm()
    
    return render(request, 'tool_app/ssl_checker.html', {
        'form': form,
        'ssl_info': ssl_info
    })


def email_validator(request):
    """Email validation tool"""
    validation_results = None
    if request.method == 'POST':
        form = EmailValidatorForm(request.POST)
        if form.is_valid():
            try:
                email_list = form.cleaned_data['email_list']
                check_mx = form.cleaned_data['check_mx']
                validation_results = validate_emails(email_list, check_mx)
            except Exception as e:
                messages.error(request, f'Error validating emails: {str(e)}')
    else:
        form = EmailValidatorForm()
    
    return render(request, 'tool_app/email_validator.html', {
        'form': form,
        'validation_results': validation_results,
        'validation_results_json': json.dumps(validation_results) if validation_results else '[]'
    })


def text_encryption(request):
    """Text encryption/decryption tool"""
    result = None
    if request.method == 'POST':
        form = TextEncryptionForm(request.POST)
        if form.is_valid():
            try:
                operation = form.cleaned_data['operation']
                algorithm = form.cleaned_data['algorithm']
                text_input = form.cleaned_data['text_input']
                key = form.cleaned_data.get('key', 3)
                result = process_text_encryption(operation, algorithm, text_input, key)
            except Exception as e:
                messages.error(request, f'Error processing text: {str(e)}')
    else:
        form = TextEncryptionForm()
    
    return render(request, 'tool_app/text_encryption.html', {
        'form': form,
        'result': result
    })


# Security Tools Helper Functions

def generate_hashes(text_input, hash_types):
    """Generate various hashes for the given text"""
    results = {}
    text_bytes = text_input.encode('utf-8')
    
    for hash_type in hash_types:
        if hash_type == 'md5':
            results['MD5'] = hashlib.md5(text_bytes).hexdigest()
        elif hash_type == 'sha1':
            results['SHA1'] = hashlib.sha1(text_bytes).hexdigest()
        elif hash_type == 'sha256':
            results['SHA256'] = hashlib.sha256(text_bytes).hexdigest()
        elif hash_type == 'sha512':
            results['SHA512'] = hashlib.sha512(text_bytes).hexdigest()
        elif hash_type == 'sha224':
            results['SHA224'] = hashlib.sha224(text_bytes).hexdigest()
        elif hash_type == 'sha384':
            results['SHA384'] = hashlib.sha384(text_bytes).hexdigest()
    
    return results


def decode_jwt_token(jwt_token):
    """Decode JWT token without verification"""
    try:
        # Split the JWT token
        parts = jwt_token.split('.')
        if len(parts) != 3:
            raise ValueError('Invalid JWT format. JWT should have 3 parts separated by dots.')
        
        header_encoded, payload_encoded, signature = parts
        
        # Add padding if needed
        header_encoded += '=' * (4 - len(header_encoded) % 4)
        payload_encoded += '=' * (4 - len(payload_encoded) % 4)
        
        # Decode header and payload
        try:
            header = json.loads(base64.urlsafe_b64decode(header_encoded))
        except Exception:
            raise ValueError('Invalid JWT header encoding')
        
        try:
            payload = json.loads(base64.urlsafe_b64decode(payload_encoded))
        except Exception:
            raise ValueError('Invalid JWT payload encoding')
        
        # Convert timestamps to readable dates
        timestamp_fields = ['exp', 'iat', 'nbf']
        for field in timestamp_fields:
            if field in payload and isinstance(payload[field], (int, float)):
                try:
                    payload[f'{field}_readable'] = datetime.fromtimestamp(payload[field], timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')
                except:
                    payload[f'{field}_readable'] = 'Invalid timestamp'
        
        return {
            'header': header,
            'payload': payload,
            'signature': signature[:20] + '...' if len(signature) > 20 else signature,
            'valid_format': True
        }
        
    except Exception as e:
        return {
            'error': str(e),
            'valid_format': False
        }


def check_ssl_certificate(domain, port=443):
    """Check SSL certificate information"""
    try:
        # Create SSL context
        context = ssl.create_default_context()
        
        # Connect to the server
        with socket.create_connection((domain, port), timeout=10) as sock:
            with context.wrap_socket(sock, server_hostname=domain) as ssock:
                cert = ssock.getpeercert()
                
                # Parse certificate information
                subject = dict(x[0] for x in cert['subject'])
                issuer = dict(x[0] for x in cert['issuer'])
                
                # Parse dates
                not_before = datetime.strptime(cert['notBefore'], '%b %d %H:%M:%S %Y %Z')
                not_after = datetime.strptime(cert['notAfter'], '%b %d %H:%M:%S %Y %Z')
                
                # Calculate days until expiry
                days_until_expiry = (not_after - datetime.now()).days
                
                # Check if expired
                is_expired = datetime.now() > not_after
                is_valid = not is_expired and datetime.now() > not_before
                
                # Get SAN (Subject Alternative Names)
                san_list = []
                if 'subjectAltName' in cert:
                    san_list = [name[1] for name in cert['subjectAltName'] if name[0] == 'DNS']
                
                return {
                    'domain': domain,
                    'port': port,
                    'subject': subject,
                    'issuer': issuer,
                    'not_before': not_before.strftime('%Y-%m-%d %H:%M:%S'),
                    'not_after': not_after.strftime('%Y-%m-%d %H:%M:%S'),
                    'days_until_expiry': days_until_expiry,
                    'is_expired': is_expired,
                    'is_valid': is_valid,
                    'san_list': san_list,
                    'version': cert.get('version', 'Unknown'),
                    'serial_number': cert.get('serialNumber', 'Unknown'),
                    'signature_algorithm': cert.get('signatureAlgorithm', 'Unknown')
                }
                
    except Exception as e:
        return {
            'error': f'SSL check failed: {str(e)}',
            'domain': domain,
            'port': port
        }


def validate_emails(email_list, check_mx=True):
    """Validate a list of email addresses"""
    results = []
    email_regex = re.compile(
        r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    )
    
    emails = [email.strip() for email in email_list.split('\n') if email.strip()]
    
    for email in emails:
        result = {
            'email': email,
            'is_valid_format': bool(email_regex.match(email)),
            'mx_valid': None,
            'mx_records': [],
            'errors': []
        }
        
        # Parse email address
        try:
            parsed_name, parsed_addr = parseaddr(email)
            if parsed_addr != email:
                result['errors'].append('Email format contains display name or invalid characters')
        except Exception:
            result['errors'].append('Failed to parse email address')
        
        # Check MX records if requested and format is valid
        if check_mx and result['is_valid_format']:
            try:
                domain = email.split('@')[1]
                mx_records = dns.resolver.resolve(domain, 'MX')
                result['mx_records'] = [str(mx) for mx in mx_records]
                result['mx_valid'] = len(result['mx_records']) > 0
            except dns.resolver.NXDOMAIN:
                result['mx_valid'] = False
                result['errors'].append('Domain does not exist')
            except dns.resolver.NoAnswer:
                result['mx_valid'] = False
                result['errors'].append('No MX records found')
            except Exception as e:
                result['mx_valid'] = False
                result['errors'].append(f'MX lookup failed: {str(e)}')
        
        results.append(result)
    
    return results


def process_text_encryption(operation, algorithm, text_input, key=3):
    """Process text encryption/decryption"""
    try:
        if algorithm == 'base64':
            if operation == 'encrypt':
                result = base64.b64encode(text_input.encode('utf-8')).decode('utf-8')
            else:  # decrypt
                result = base64.b64decode(text_input.encode('utf-8')).decode('utf-8')
                
        elif algorithm == 'caesar':
            if operation == 'encrypt':
                result = caesar_cipher(text_input, key)
            else:  # decrypt
                result = caesar_cipher(text_input, -key)
                
        elif algorithm == 'rot13':
            # ROT13 is its own inverse
            result = text_input.translate(str.maketrans(
                'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz',
                'NOPQRSTUVWXYZABCDEFGHIJKLMnopqrstuvwxyzabcdefghijklm'
            ))
        
        else:
            raise ValueError(f'Unsupported algorithm: {algorithm}')
        
        return {
            'success': True,
            'result': result,
            'operation': operation,
            'algorithm': algorithm,
            'original_text': text_input
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'operation': operation,
            'algorithm': algorithm
        }


def caesar_cipher(text, shift):
    """Apply Caesar cipher with given shift"""
    result = []
    
    for char in text:
        if char.isalpha():
            # Determine if uppercase or lowercase
            is_upper = char.isupper()
            char = char.lower()
            
            # Apply shift
            shifted = chr((ord(char) - ord('a') + shift) % 26 + ord('a'))
            
            # Restore case
            if is_upper:
                shifted = shifted.upper()
            
            result.append(shifted)
        else:
            # Keep non-alphabetic characters unchanged
            result.append(char)
    
    return ''.join(result)


# Fun & Engagement Tools Views

def meme_generator(request):
    """Meme generator tool"""
    meme_result = None
    if request.method == 'POST':
        form = MemeGeneratorForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                image_file = form.cleaned_data['image_file']
                top_text = form.cleaned_data['top_text']
                bottom_text = form.cleaned_data['bottom_text']
                font_size = form.cleaned_data['font_size']
                meme_result = create_meme(image_file, top_text, bottom_text, font_size)
            except Exception as e:
                messages.error(request, f'Error creating meme: {str(e)}')
    else:
        form = MemeGeneratorForm()
    
    return render(request, 'tool_app/meme_generator.html', {
        'form': form,
        'meme_result': meme_result
    })


def emoji_translator(request):
    """Emoji translator tool"""
    translation_result = None
    if request.method == 'POST':
        form = EmojiTranslatorForm(request.POST)
        if form.is_valid():
            try:
                text_input = form.cleaned_data['text_input']
                translation_mode = form.cleaned_data['translation_mode']
                translation_result = translate_to_emoji(text_input, translation_mode)
            except Exception as e:
                messages.error(request, f'Error translating text: {str(e)}')
    else:
        form = EmojiTranslatorForm()
    
    return render(request, 'tool_app/emoji_translator.html', {
        'form': form,
        'translation_result': translation_result
    })


def random_quote(request):
    """Random quote generator tool"""
    quotes_result = None
    if request.method == 'POST':
        form = RandomQuoteForm(request.POST)
        if form.is_valid():
            try:
                category = form.cleaned_data['category']
                count = form.cleaned_data['count']
                quotes_result = generate_random_quotes(category, count)
            except Exception as e:
                messages.error(request, f'Error generating quotes: {str(e)}')
    else:
        form = RandomQuoteForm()
    
    return render(request, 'tool_app/random_quote.html', {
        'form': form,
        'quotes_result': quotes_result
    })


def random_name(request):
    """Random name generator tool"""
    names_result = None
    if request.method == 'POST':
        form = RandomNameForm(request.POST)
        if form.is_valid():
            try:
                name_type = form.cleaned_data['name_type']
                gender = form.cleaned_data['gender']
                count = form.cleaned_data['count']
                names_result = generate_random_names(name_type, gender, count)
            except Exception as e:
                messages.error(request, f'Error generating names: {str(e)}')
    else:
        form = RandomNameForm()
    
    return render(request, 'tool_app/random_name.html', {
        'form': form,
        'names_result': names_result
    })


def unit_converter(request):
    """Unit converter tool"""
    conversion_result = None
    if request.method == 'POST':
        form = UnitConverterForm(request.POST)
        if form.is_valid():
            try:
                conversion_type = form.cleaned_data['conversion_type']
                value = form.cleaned_data['value']
                from_unit = form.cleaned_data['from_unit']
                to_unit = form.cleaned_data['to_unit']
                conversion_result = convert_units(conversion_type, value, from_unit, to_unit)
            except Exception as e:
                messages.error(request, f'Error converting units: {str(e)}')
    else:
        form = UnitConverterForm()
    
    return render(request, 'tool_app/unit_converter.html', {
        'form': form,
        'conversion_result': conversion_result
    })


# Fun Tools Helper Functions

def create_meme(image_file, top_text, bottom_text, font_size):
    """Create a meme from image and text"""
    try:
        # Open and process the image
        image = Image.open(image_file)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Get image dimensions
        width, height = image.size
        
        # Create drawing context
        draw = ImageDraw.Draw(image)
        
        # Try to use a better font, fall back to default
        try:
            font = ImageFont.truetype("arial.ttf", font_size)
        except:
            try:
                font = ImageFont.truetype("DejaVuSans-Bold.ttf", font_size)
            except:
                font = ImageFont.load_default()
        
        # Add top text
        if top_text:
            # Wrap text if needed
            wrapped_text = textwrap.fill(top_text.upper(), width=20)
            lines = wrapped_text.split('\n')
            
            for i, line in enumerate(lines):
                # Get text size
                bbox = draw.textbbox((0, 0), line, font=font)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]
                
                # Calculate position (centered, near top)
                x = (width - text_width) // 2
                y = 10 + (i * (text_height + 5))
                
                # Draw text with outline
                outline_width = 2
                for dx in range(-outline_width, outline_width + 1):
                    for dy in range(-outline_width, outline_width + 1):
                        draw.text((x + dx, y + dy), line, font=font, fill='black')
                draw.text((x, y), line, font=font, fill='white')
        
        # Add bottom text
        if bottom_text:
            # Wrap text if needed
            wrapped_text = textwrap.fill(bottom_text.upper(), width=20)
            lines = wrapped_text.split('\n')
            
            for i, line in enumerate(reversed(lines)):
                # Get text size
                bbox = draw.textbbox((0, 0), line, font=font)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]
                
                # Calculate position (centered, near bottom)
                x = (width - text_width) // 2
                y = height - 10 - (text_height * (i + 1)) - (i * 5)
                
                # Draw text with outline
                outline_width = 2
                for dx in range(-outline_width, outline_width + 1):
                    for dy in range(-outline_width, outline_width + 1):
                        draw.text((x + dx, y + dy), line, font=font, fill='black')
                draw.text((x, y), line, font=font, fill='white')
        
        # Save to buffer
        buffer = BytesIO()
        image.save(buffer, format='JPEG', quality=95)
        buffer.seek(0)
        
        # Encode to base64 for display
        import base64
        image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        return {
            'success': True,
            'image_data': image_base64,
            'top_text': top_text,
            'bottom_text': bottom_text,
            'font_size': font_size
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def translate_to_emoji(text_input, translation_mode):
    """Translate text to emoji"""
    # Basic emoji mapping
    emoji_map = {
        'happy': '😊', 'sad': '😢', 'love': '❤️', 'heart': '💖', 'fire': '🔥',
        'water': '💧', 'sun': '☀️', 'moon': '🌙', 'star': '⭐', 'tree': '🌳',
        'flower': '🌸', 'cat': '🐱', 'dog': '🐶', 'bird': '🐦', 'fish': '🐟',
        'car': '🚗', 'plane': '✈️', 'train': '🚂', 'bike': '🚲', 'house': '🏠',
        'food': '🍔', 'pizza': '🍕', 'cake': '🎂', 'coffee': '☕', 'beer': '🍺',
        'music': '🎵', 'book': '📚', 'phone': '📱', 'computer': '💻', 'game': '🎮',
        'money': '💰', 'work': '💼', 'time': '⏰', 'sleep': '💤', 'party': '🎉',
        'good': '👍', 'bad': '👎', 'ok': '👌', 'yes': '✅', 'no': '❌',
        'hot': '🥵', 'cold': '🥶', 'cool': '😎', 'wow': '😮', 'laugh': '😂',
        'cry': '😭', 'angry': '😠', 'think': '🤔', 'idea': '💡', 'strong': '💪'
    }
    
    words = text_input.split()
    result_words = []
    
    for word in words:
        word_lower = word.lower().strip('.,!?;:')
        emoji = emoji_map.get(word_lower, '')
        
        if translation_mode == 'replace' and emoji:
            result_words.append(emoji)
        elif translation_mode == 'append' and emoji:
            result_words.append(f'{word} {emoji}')
        elif translation_mode == 'prepend' and emoji:
            result_words.append(f'{emoji} {word}')
        else:
            result_words.append(word)
    
    return {
        'original_text': text_input,
        'translated_text': ' '.join(result_words),
        'translation_mode': translation_mode,
        'emojis_found': len([w for w in words if w.lower().strip('.,!?;:') in emoji_map])
    }


def generate_random_quotes(category, count):
    """Generate random quotes"""
    quotes_db = {
        'inspirational': [
            {'quote': 'The only way to do great work is to love what you do.', 'author': 'Steve Jobs'},
            {'quote': 'Innovation distinguishes between a leader and a follower.', 'author': 'Steve Jobs'},
            {'quote': 'Life is what happens to you while you\'re busy making other plans.', 'author': 'John Lennon'},
            {'quote': 'The future belongs to those who believe in the beauty of their dreams.', 'author': 'Eleanor Roosevelt'},
            {'quote': 'It is during our darkest moments that we must focus to see the light.', 'author': 'Aristotle'},
        ],
        'motivational': [
            {'quote': 'Success is not final, failure is not fatal: it is the courage to continue that counts.', 'author': 'Winston Churchill'},
            {'quote': 'The only impossible journey is the one you never begin.', 'author': 'Tony Robbins'},
            {'quote': 'Don\'t watch the clock; do what it does. Keep going.', 'author': 'Sam Levenson'},
            {'quote': 'The way to get started is to quit talking and begin doing.', 'author': 'Walt Disney'},
            {'quote': 'Believe you can and you\'re halfway there.', 'author': 'Theodore Roosevelt'},
        ],
        'funny': [
            {'quote': 'I have not failed. I\'ve just found 10,000 ways that won\'t work.', 'author': 'Thomas Edison'},
            {'quote': 'The only thing worse than being talked about is not being talked about.', 'author': 'Oscar Wilde'},
            {'quote': 'I can resist everything except temptation.', 'author': 'Oscar Wilde'},
            {'quote': 'Common sense is not so common.', 'author': 'Voltaire'},
            {'quote': 'The trouble with having an open mind is that people keep coming along and trying to put things in it.', 'author': 'Terry Pratchett'},
        ],
        'love': [
            {'quote': 'Being deeply loved by someone gives you strength, while loving someone deeply gives you courage.', 'author': 'Lao Tzu'},
            {'quote': 'The best thing to hold onto in life is each other.', 'author': 'Audrey Hepburn'},
            {'quote': 'Love is composed of a single soul inhabiting two bodies.', 'author': 'Aristotle'},
            {'quote': 'Where there is love there is life.', 'author': 'Mahatma Gandhi'},
            {'quote': 'Love all, trust a few, do wrong to none.', 'author': 'William Shakespeare'},
        ],
        'life': [
            {'quote': 'In the end, we will remember not the words of our enemies, but the silence of our friends.', 'author': 'Martin Luther King Jr.'},
            {'quote': 'Life is really simple, but we insist on making it complicated.', 'author': 'Confucius'},
            {'quote': 'The purpose of our lives is to be happy.', 'author': 'Dalai Lama'},
            {'quote': 'Life is 10% what happens to you and 90% how you react to it.', 'author': 'Charles R. Swindoll'},
            {'quote': 'Be yourself; everyone else is already taken.', 'author': 'Oscar Wilde'},
        ],
        'success': [
            {'quote': 'Success is not the key to happiness. Happiness is the key to success.', 'author': 'Albert Schweitzer'},
            {'quote': 'Success is walking from failure to failure with no loss of enthusiasm.', 'author': 'Winston Churchill'},
            {'quote': 'The road to success and the road to failure are almost exactly the same.', 'author': 'Colin R. Davis'},
            {'quote': 'Success is not in what you have, but who you are.', 'author': 'Bo Bennett'},
            {'quote': 'Don\'t be afraid to give up the good to go for the great.', 'author': 'John D. Rockefeller'},
        ],
        'wisdom': [
            {'quote': 'The only true wisdom is in knowing you know nothing.', 'author': 'Socrates'},
            {'quote': 'Yesterday is history, tomorrow is a mystery, today is a gift.', 'author': 'Eleanor Roosevelt'},
            {'quote': 'The journey of a thousand miles begins with one step.', 'author': 'Lao Tzu'},
            {'quote': 'It does not matter how slowly you go as long as you do not stop.', 'author': 'Confucius'},
            {'quote': 'A wise man can learn more from a foolish question than a fool can learn from a wise answer.', 'author': 'Bruce Lee'},
        ]
    }
    
    if category == 'all':
        all_quotes = []
        for cat_quotes in quotes_db.values():
            all_quotes.extend(cat_quotes)
        selected_quotes = random.sample(all_quotes, min(count, len(all_quotes)))
    else:
        cat_quotes = quotes_db.get(category, [])
        selected_quotes = random.sample(cat_quotes, min(count, len(cat_quotes)))
    
    return {
        'quotes': selected_quotes,
        'category': category,
        'count': len(selected_quotes)
    }


def generate_random_names(name_type, gender, count):
    """Generate random names"""
    # Sample name data
    male_first = ['James', 'John', 'Robert', 'Michael', 'William', 'David', 'Richard', 'Joseph', 'Thomas', 'Christopher',
                  'Charles', 'Daniel', 'Matthew', 'Anthony', 'Mark', 'Donald', 'Steven', 'Paul', 'Andrew', 'Kenneth']
    
    female_first = ['Mary', 'Patricia', 'Jennifer', 'Linda', 'Elizabeth', 'Barbara', 'Susan', 'Jessica', 'Sarah', 'Karen',
                    'Lisa', 'Nancy', 'Betty', 'Helen', 'Sandra', 'Donna', 'Carol', 'Ruth', 'Sharon', 'Michelle']
    
    last_names = ['Smith', 'Johnson', 'Williams', 'Brown', 'Jones', 'Garcia', 'Miller', 'Davis', 'Rodriguez', 'Martinez',
                  'Hernandez', 'Lopez', 'Gonzalez', 'Wilson', 'Anderson', 'Thomas', 'Taylor', 'Moore', 'Jackson', 'Martin']
    
    results = []
    
    for _ in range(count):
        if name_type == 'username':
            # Generate username
            adjectives = ['Cool', 'Super', 'Mega', 'Ultra', 'Pro', 'Epic', 'Awesome', 'Amazing', 'Fantastic', 'Incredible']
            nouns = ['Gamer', 'Ninja', 'Warrior', 'Master', 'Hero', 'Legend', 'Champion', 'Star', 'Phoenix', 'Dragon']
            username = f'{random.choice(adjectives)}{random.choice(nouns)}{random.randint(100, 999)}'
            results.append(username)
        
        elif name_type == 'first':
            if gender == 'male':
                results.append(random.choice(male_first))
            elif gender == 'female':
                results.append(random.choice(female_first))
            else:
                results.append(random.choice(male_first + female_first))
        
        elif name_type == 'last':
            results.append(random.choice(last_names))
        
        else:  # full name
            if gender == 'male':
                first = random.choice(male_first)
            elif gender == 'female':
                first = random.choice(female_first)
            else:
                first = random.choice(male_first + female_first)
            
            last = random.choice(last_names)
            results.append(f'{first} {last}')
    
    return {
        'names': results,
        'name_type': name_type,
        'gender': gender,
        'count': len(results)
    }


def convert_units(conversion_type, value, from_unit, to_unit):
    """Convert between different units"""
    # Unit conversion rates (to base unit)
    conversions = {
        'length': {
            'mm': 0.001, 'cm': 0.01, 'm': 1, 'km': 1000,
            'in': 0.0254, 'ft': 0.3048, 'yd': 0.9144, 'mi': 1609.34
        },
        'weight': {
            'mg': 0.000001, 'g': 0.001, 'kg': 1, 't': 1000,
            'oz': 0.0283495, 'lb': 0.453592, 'st': 6.35029
        },
        'volume': {
            'ml': 0.001, 'l': 1, 'gal': 3.78541, 'qt': 0.946353,
            'pt': 0.473176, 'cup': 0.236588, 'fl oz': 0.0295735
        },
        'area': {
            'mm²': 0.000001, 'cm²': 0.0001, 'm²': 1, 'km²': 1000000,
            'in²': 0.00064516, 'ft²': 0.092903, 'yd²': 0.836127, 'acre': 4046.86
        },
        'speed': {
            'm/s': 1, 'km/h': 0.277778, 'mph': 0.44704, 'ft/s': 0.3048, 'knot': 0.514444
        }
    }
    
    try:
        if conversion_type == 'temperature':
            # Special handling for temperature
            result = convert_temperature(value, from_unit, to_unit)
        else:
            # Standard conversion through base unit
            unit_rates = conversions[conversion_type]
            
            # Convert to base unit
            base_value = value * unit_rates[from_unit]
            
            # Convert from base unit to target
            result = base_value / unit_rates[to_unit]
        
        return {
            'success': True,
            'original_value': value,
            'original_unit': from_unit,
            'converted_value': round(result, 6),
            'converted_unit': to_unit,
            'conversion_type': conversion_type,
            'formula': f'{value} {from_unit} = {round(result, 6)} {to_unit}'
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'conversion_type': conversion_type
        }


def convert_temperature(value, from_unit, to_unit):
    """Convert temperature between Celsius, Fahrenheit, and Kelvin"""
    # Convert to Celsius first
    if from_unit == 'F':
        celsius = (value - 32) * 5/9
    elif from_unit == 'K':
        celsius = value - 273.15
    else:  # from_unit == 'C'
        celsius = value
    
    # Convert from Celsius to target
    if to_unit == 'F':
        return celsius * 9/5 + 32
    elif to_unit == 'K':
        return celsius + 273.15
    else:  # to_unit == 'C'
        return celsius


# Audio/Video Tools Views

def audio_converter(request):
    """Audio format converter tool"""
    conversion_result = None
    if request.method == 'POST':
        form = AudioConverterForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                audio_file = form.cleaned_data['audio_file']
                target_format = form.cleaned_data['target_format']
                quality = form.cleaned_data['quality']
                conversion_result = convert_audio_format(audio_file, target_format, quality)
            except Exception as e:
                messages.error(request, f'Error converting audio: {str(e)}')
    else:
        form = AudioConverterForm()
    
    return render(request, 'tool_app/audio_converter.html', {
        'form': form,
        'conversion_result': conversion_result
    })


def audio_speed_changer(request):
    """Audio speed changer tool"""
    processing_result = None
    if request.method == 'POST':
        form = AudioSpeedChangerForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                audio_file = form.cleaned_data['audio_file']
                speed_multiplier = form.cleaned_data['speed_multiplier']
                preserve_pitch = form.cleaned_data['preserve_pitch']
                output_format = form.cleaned_data['output_format']
                processing_result = change_audio_speed(audio_file, speed_multiplier, preserve_pitch, output_format)
            except Exception as e:
                messages.error(request, f'Error changing audio speed: {str(e)}')
    else:
        form = AudioSpeedChangerForm()
    
    return render(request, 'tool_app/audio_speed_changer.html', {
        'form': form,
        'processing_result': processing_result
    })


def video_to_audio(request):
    """Video to audio extractor tool"""
    extraction_result = None
    if request.method == 'POST':
        form = VideoToAudioForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                video_file = form.cleaned_data['video_file']
                audio_format = form.cleaned_data['audio_format']
                audio_quality = form.cleaned_data['audio_quality']
                start_time = form.cleaned_data.get('start_time')
                end_time = form.cleaned_data.get('end_time')
                extraction_result = extract_audio_from_video(
                    video_file, audio_format, audio_quality, start_time, end_time
                )
            except Exception as e:
                messages.error(request, f'Error extracting audio: {str(e)}')
    else:
        form = VideoToAudioForm()
    
    return render(request, 'tool_app/video_to_audio.html', {
        'form': form,
        'extraction_result': extraction_result
    })


# Audio/Video Processing Helper Functions

def convert_audio_format(audio_file, target_format, quality):
    """Convert audio file to different format"""
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{audio_file.name.split(".")[-1]}') as temp_input:
            # Write uploaded file to temp file
            for chunk in audio_file.chunks():
                temp_input.write(chunk)
            temp_input_path = temp_input.name
        
        # Load audio with pydub
        audio = AudioSegment.from_file(temp_input_path)
        
        # Get audio information
        original_info = {
            'duration': len(audio) / 1000.0,  # seconds
            'channels': audio.channels,
            'sample_rate': audio.frame_rate,
            'bitrate': audio.frame_rate * audio.frame_width * 8 * audio.channels // 1000 if audio.frame_width else 'Unknown'
        }
        
        # Create output temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{target_format}') as temp_output:
            temp_output_path = temp_output.name
        
        # Export with specified quality
        export_params = {'format': target_format}
        if target_format in ['mp3', 'aac', 'ogg']:
            export_params['bitrate'] = f'{quality}k'
        
        audio.export(temp_output_path, **export_params)
        
        # Read the converted file
        with open(temp_output_path, 'rb') as f:
            converted_data = f.read()
        
        # Get converted file info
        converted_audio = AudioSegment.from_file(temp_output_path)
        converted_info = {
            'duration': len(converted_audio) / 1000.0,
            'channels': converted_audio.channels,
            'sample_rate': converted_audio.frame_rate,
            'size': len(converted_data)
        }
        
        # Clean up temp files
        os.unlink(temp_input_path)
        os.unlink(temp_output_path)
        
        # Encode to base64 for download
        import base64
        audio_base64 = base64.b64encode(converted_data).decode('utf-8')
        
        return {
            'success': True,
            'audio_data': audio_base64,
            'original_format': audio_file.name.split('.')[-1].upper(),
            'target_format': target_format.upper(),
            'quality': quality,
            'filename': f"converted.{target_format}",
            'original_info': original_info,
            'converted_info': converted_info
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def change_audio_speed(audio_file, speed_multiplier, preserve_pitch, output_format):
    """Change audio playback speed"""
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{audio_file.name.split(".")[-1]}') as temp_input:
            for chunk in audio_file.chunks():
                temp_input.write(chunk)
            temp_input_path = temp_input.name
        
        # Load audio
        audio = AudioSegment.from_file(temp_input_path)
        
        # Get original info
        original_duration = len(audio) / 1000.0
        
        # Change speed
        if preserve_pitch:
            # Use speedup method that preserves pitch (simple approach)
            # Note: This is a basic implementation. For professional use, 
            # consider using librosa or other advanced audio processing libraries
            if speed_multiplier > 1:
                # Speed up: remove samples
                audio = audio[::int(speed_multiplier)]
            else:
                # Slow down: interpolate (basic approach)
                audio = audio._spawn(audio.raw_data, overrides={'frame_rate': int(audio.frame_rate * speed_multiplier)})
                audio = audio.set_frame_rate(audio.frame_rate)
        else:
            # Simple speed change (changes pitch)
            new_sample_rate = int(audio.frame_rate * speed_multiplier)
            audio = audio._spawn(audio.raw_data, overrides={'frame_rate': new_sample_rate})
            audio = audio.set_frame_rate(audio.frame_rate)
        
        # Create output temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{output_format}') as temp_output:
            temp_output_path = temp_output.name
        
        # Export processed audio
        audio.export(temp_output_path, format=output_format)
        
        # Read the processed file
        with open(temp_output_path, 'rb') as f:
            processed_data = f.read()
        
        # Clean up
        os.unlink(temp_input_path)
        os.unlink(temp_output_path)
        
        # Calculate new duration
        new_duration = original_duration / speed_multiplier
        
        # Encode to base64
        import base64
        audio_base64 = base64.b64encode(processed_data).decode('utf-8')
        
        return {
            'success': True,
            'audio_data': audio_base64,
            'filename': f"speed_changed.{output_format}",
            'speed_multiplier': speed_multiplier,
            'preserve_pitch': preserve_pitch,
            'output_format': output_format.upper(),
            'original_duration': original_duration,
            'new_duration': new_duration,
            'size': len(processed_data)
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def extract_audio_from_video(video_file, audio_format, audio_quality, start_time=None, end_time=None):
    """Extract audio from video file"""
    try:
        # Check if VideoFileClip is available
        if VideoFileClip is None:
            return {
                'success': False,
                'error': 'MoviePy library is not properly installed. Please install it with: pip install moviepy'
            }
            
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{video_file.name.split(".")[-1]}') as temp_input:
            for chunk in video_file.chunks():
                temp_input.write(chunk)
            temp_input_path = temp_input.name
        
        # Load video with moviepy
        video_clip = VideoFileClip(temp_input_path)
        
        # Get video information
        video_info = {
            'duration': video_clip.duration,
            'fps': video_clip.fps,
            'size': video_clip.size,
            'has_audio': video_clip.audio is not None
        }
        
        if not video_clip.audio:
            video_clip.close()
            os.unlink(temp_input_path)
            return {
                'success': False,
                'error': 'This video file does not contain audio'
            }
        
        # Extract audio
        audio_clip = video_clip.audio
        
        # Apply time range if specified
        if start_time or end_time:
            start_seconds = parse_time_to_seconds(start_time) if start_time else 0
            end_seconds = parse_time_to_seconds(end_time) if end_time else video_clip.duration
            
            if start_seconds < end_seconds and start_seconds >= 0 and end_seconds <= video_clip.duration:
                audio_clip = audio_clip.subclip(start_seconds, end_seconds)
            else:
                video_clip.close()
                os.unlink(temp_input_path)
                return {
                    'success': False,
                    'error': 'Invalid time range specified'
                }
        
        # Create output temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{audio_format}') as temp_output:
            temp_output_path = temp_output.name
        
        # Export audio with quality settings
        audio_params = {}
        if audio_format == 'mp3':
            audio_params['bitrate'] = f'{audio_quality}k'
        elif audio_format == 'aac':
            audio_params['bitrate'] = f'{audio_quality}k'
        
        audio_clip.write_audiofile(temp_output_path, **audio_params, verbose=False, logger=None)
        
        # Read the extracted audio file
        with open(temp_output_path, 'rb') as f:
            audio_data = f.read()
        
        # Get audio information
        audio_info = {
            'duration': audio_clip.duration,
            'size': len(audio_data),
            'format': audio_format.upper(),
            'quality': f'{audio_quality} kbps'
        }
        
        # Clean up
        video_clip.close()
        audio_clip.close()
        os.unlink(temp_input_path)
        os.unlink(temp_output_path)
        
        # Encode to base64
        import base64
        audio_base64 = base64.b64encode(audio_data).decode('utf-8')
        
        return {
            'success': True,
            'audio_data': audio_base64,
            'filename': f"extracted_audio.{audio_format}",
            'video_info': video_info,
            'audio_info': audio_info,
            'start_time': start_time,
            'end_time': end_time
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def parse_time_to_seconds(time_str):
    """Parse time string (MM:SS or HH:MM:SS) to seconds"""
    if not time_str:
        return 0
    
    parts = time_str.split(':')
    if len(parts) == 2:  # MM:SS
        minutes, seconds = map(int, parts)
        return minutes * 60 + seconds
    elif len(parts) == 3:  # HH:MM:SS
        hours, minutes, seconds = map(int, parts)
        return hours * 3600 + minutes * 60 + seconds
    else:
        raise ValueError('Invalid time format')


# Static Page Views
def about(request):
    """About us page"""
    return render(request, 'tool_app/about.html')


def contact(request):
    """Contact us page"""
    if request.method == 'POST':
        # Handle contact form submission
        # In a real implementation, you would process the form and send email
        messages.success(request, 'Thank you for your message! We will get back to you soon.')
        return redirect('tool_app:contact')
    
    return render(request, 'tool_app/contact.html')


def privacy_policy(request):
    """Privacy policy page"""
    return render(request, 'tool_app/privacy_policy.html')


def terms(request):
    """Terms and conditions page"""
    return render(request, 'tool_app/terms.html')


def disclaimer(request):
    """Disclaimer page"""
    return render(request, 'tool_app/disclaimer.html')


def blog(request):
    """Blog page"""
    return render(request, 'tool_app/blog.html')
